import os
from datetime import datetime

# DOCX generation
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# PDF generation
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image


OUTPUT_DIR = "/workspace/output"


def ensure_output_dir() -> None:
    if not os.path.isdir(OUTPUT_DIR):
        os.makedirs(OUTPUT_DIR, exist_ok=True)


PROFILE_IMAGE_CANDIDATES = [
    "/workspace/profile.jpg",
    "/workspace/profile.jpeg",
    "/workspace/profile.png",
]


def find_profile_image() -> str | None:
    for path in PROFILE_IMAGE_CANDIDATES:
        if os.path.exists(path):
            return path
    # Also accept image passed via environment variable
    env_path = os.environ.get("PROFILE_IMAGE_PATH")
    if env_path and os.path.exists(env_path):
        return env_path
    return None


def build_content() -> dict:
    return {
        "name": "VIKRAM SINGH GILL",
        "location": "Jaipur, Rajasthan, India",
        "phone": "9571041410",
        "email": "vkmgudha@gmail.com",
        "objective": (
            "Motivated and dedicated hospitality professional seeking a position in the "
            "Food & Beverage Department where I can apply my service skills, customer care "
            "abilities, and eagerness to learn in order to contribute to guest satisfaction "
            "and organizational success."
        ),
        "education": [
            ("Diploma in Hotel Management", "Konark College of Hotel Management, Sikar (2022)"),
            ("B.A.", "Pandit Deendayal Upadhyaya Shekhawati University, Sikar (2021)"),
            ("Senior Secondary (Class 12)", "RBSE, Ajmer (2016)"),
            ("Secondary (Class 10)", "RBSE, Ajmer (2014)"),
        ],
        "experience": [
            {
                "title": "Guest Service Attendant – The Grand Dragon, Leh-Ladakh",
                "dates": "July 2022 – Sept 2024",
                "bullets": [
                    "Assisted guests with check-in/check-out and dining services.",
                    "Ensured high standards of customer service in a luxury environment.",
                    "Maintained professionalism while handling diverse guest needs.",
                ],
            },
            {
                "title": "Service Staff – Nothing Before Coffee (NBC), Jaipur",
                "dates": "Feb 2025 – Sept 2025",
                "bullets": [
                    "Delivered efficient and friendly service in a fast-paced café environment.",
                    "Managed customer orders and ensured quality service.",
                    "Built strong rapport with regular customers.",
                ],
            },
        ],
        "skills": [
            "Customer Service & Guest Relations",
            "Food & Beverage Service",
            "Teamwork & Time Management",
            "Basic Computer Knowledge",
            "Languages: English & Hindi",
        ],
        "personal": {
            "Date of Birth": "24 July 2000",
            "Marital Status": "Unmarried",
            "Nationality": "Indian",
        },
        "place": "Jaipur",
    }


def add_heading(document: Document, text: str) -> None:
    p = document.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    p.space_before = Pt(8)
    p.space_after = Pt(2)


def generate_docx(content: dict, image_path: str | None) -> str:
    document = Document()

    # Set default font
    style = document.styles['Normal']
    style.font.name = 'Calibri'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    style.font.size = Pt(10.5)

    # Header with name and contact
    header_table = document.add_table(rows=1, cols=2)
    header_table.autofit = True
    left_cell, right_cell = header_table.rows[0].cells

    # Name and contacts
    name_para = left_cell.paragraphs[0]
    name_run = name_para.add_run(content["name"])
    name_run.bold = True
    name_run.font.size = Pt(18)
    name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    contact_para = left_cell.add_paragraph(
        f"📍 {content['location']} | 📞 {content['phone']} | ✉️ {content['email']}"
    )
    contact_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if image_path and os.path.exists(image_path):
        # Add profile image on the right
        p = right_cell.paragraphs[0]
        run = p.add_run()
        try:
            run.add_picture(image_path, width=Inches(1.2))
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        except Exception:
            pass

    # Objective
    add_heading(document, "Career Objective")
    document.add_paragraph(content["objective"])    

    # Education
    add_heading(document, "Education")
    edu_table = document.add_table(rows=1, cols=2)
    hdr_cells = edu_table.rows[0].cells
    hdr_cells[0].text = "Qualification"
    hdr_cells[1].text = "Institute / Year"
    for degree, inst in content["education"]:
        row_cells = edu_table.add_row().cells
        row_cells[0].text = degree
        row_cells[1].text = inst

    # Experience
    add_heading(document, "Work Experience")
    for exp in content["experience"]:
        p = document.add_paragraph()
        r = p.add_run(f"{exp['title']} | {exp['dates']}")
        r.bold = True
        for bullet in exp["bullets"]:
            document.add_paragraph(bullet, style='List Bullet')

    # Skills
    add_heading(document, "Skills")
    for skill in content["skills"]:
        document.add_paragraph(skill, style='List Bullet')

    # Personal Details
    add_heading(document, "Personal Details")
    for k, v in content["personal"].items():
        document.add_paragraph(f"{k}: {v}")

    # Declaration
    add_heading(document, "Declaration")
    document.add_paragraph(
        "I hereby declare that the above information is true to my knowledge."
    )
    document.add_paragraph(f"Place: {content['place']}")
    document.add_paragraph(f"Date: {datetime.now().strftime('%d-%m-%Y')}")
    document.add_paragraph("(Vikram Singh Gill)")

    ensure_output_dir()
    out_path = os.path.join(OUTPUT_DIR, "Vikram_Singh_Gill_CV.docx")
    document.save(out_path)
    return out_path


def generate_pdf(content: dict, image_path: str | None) -> str:
    ensure_output_dir()
    out_path = os.path.join(OUTPUT_DIR, "Vikram_Singh_Gill_CV.pdf")

    doc = SimpleDocTemplate(
        out_path,
        pagesize=A4,
        leftMargin=36,
        rightMargin=36,
        topMargin=36,
        bottomMargin=36,
        title=content["name"],
        author=content["name"],
    )

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="HeadingSmall", fontSize=11, leading=14, spaceBefore=12, spaceAfter=4, textColor=colors.black, alignment=TA_LEFT))

    elements = []

    # Header
    header_data = [
        [
            Paragraph(f"<b>{content['name']}</b>", styles['Title']),
            Image(image_path, width=80, height=100) if image_path and os.path.exists(image_path) else "",
        ],
        [
            Paragraph(
                f"📍 {content['location']}  |  📞 {content['phone']}  |  ✉️ {content['email']}",
                styles['Normal'],
            ),
            "",
        ],
    ]
    header_table = Table(header_data, colWidths=[400, 100])
    header_table.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
    ]))
    elements.append(header_table)
    elements.append(Spacer(1, 8))

    # Objective
    elements.append(Paragraph("<b>Career Objective</b>", styles['HeadingSmall']))
    elements.append(Paragraph(content['objective'], styles['Normal']))

    # Education
    elements.append(Paragraph("<b>Education</b>", styles['HeadingSmall']))
    edu_rows = [["Qualification", "Institute / Year"]]
    edu_rows.extend(list(content['education']))
    edu_table = Table(edu_rows, colWidths=[220, 280])
    edu_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
        ('INNERGRID', (0, 0), (-1, -1), 0.25, colors.grey),
        ('BOX', (0, 0), (-1, -1), 0.25, colors.grey),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
    ]))
    elements.append(edu_table)
    elements.append(Spacer(1, 6))

    # Experience
    elements.append(Paragraph("<b>Work Experience</b>", styles['HeadingSmall']))
    for exp in content['experience']:
        elements.append(Paragraph(f"<b>{exp['title']}</b> | {exp['dates']}", styles['Normal']))
        for bullet in exp['bullets']:
            elements.append(Paragraph(f"• {bullet}", styles['Normal']))
        elements.append(Spacer(1, 4))

    # Skills
    elements.append(Paragraph("<b>Skills</b>", styles['HeadingSmall']))
    for skill in content['skills']:
        elements.append(Paragraph(f"• {skill}", styles['Normal']))

    # Personal Details
    elements.append(Paragraph("<b>Personal Details</b>", styles['HeadingSmall']))
    for k, v in content['personal'].items():
        elements.append(Paragraph(f"{k}: {v}", styles['Normal']))

    # Declaration
    elements.append(Paragraph("<b>Declaration</b>", styles['HeadingSmall']))
    elements.append(Paragraph(
        "I hereby declare that the above information is true to my knowledge.",
        styles['Normal'],
    ))
    elements.append(Paragraph(f"Place: {content['place']}", styles['Normal']))
    elements.append(Paragraph(f"Date: {datetime.now().strftime('%d-%m-%Y')}", styles['Normal']))
    elements.append(Paragraph("(Vikram Singh Gill)", styles['Normal']))

    doc.build(elements)
    return out_path


def main() -> None:
    ensure_output_dir()
    image_path = find_profile_image()
    content = build_content()
    docx_path = generate_docx(content, image_path)
    pdf_path = generate_pdf(content, image_path)
    print(f"Generated DOCX: {docx_path}")
    print(f"Generated PDF:  {pdf_path}")


if __name__ == "__main__":
    main()

